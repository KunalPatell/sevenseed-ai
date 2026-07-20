"""Markdown -> professional .docx builder for Sevenseed's Business Document Studio.

Ported/genericized from a proven BA document-automation pipeline: cover page,
branded header/footer with page numbers, styled headings, Markdown tables, code
blocks, blockquotes and nested lists. python-docx is imported at module load, so
this module is imported lazily (only when a DOCX is actually requested) — the app
runs fine without python-docx installed.
"""
import io
import os
import re
import datetime
import logging

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

logger = logging.getLogger(__name__)

BRAND = "Sevenseed Venture Studio"
LOGO_PATH = os.environ.get("SEVENSEED_LOGO_PATH", "")  # optional image; text wordmark used if absent
ACCENT = "1A3C6B"  # navy

_DOC_FULLNAMES = {
    "BRD": "Business Requirements Document (BRD)",
    "SRS": "Software Requirements Specification (SRS)",
    "FRS": "Functional Requirement Specification (FRS)",
    "PRD": "Product Requirements Document (PRD)",
    "CHARTER": "Project Charter",
    "BUSINESS_PLAN": "Business Plan",
    "SOW": "Statement of Work (SOW)",
    "GDD": "Game Design Document (GDD)",
}


def _set_cell_bg(cell, hex_color: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _add_border(para, position: str = "bottom", color: str = ACCENT,
                sz: str = "6", space: str = "4") -> None:
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bd = OxmlElement(f"w:{position}")
    bd.set(qn("w:val"), "single")
    bd.set(qn("w:sz"), sz)
    bd.set(qn("w:space"), space)
    bd.set(qn("w:color"), color)
    pBdr.append(bd)
    pPr.append(pBdr)


def _field_run(run, field_name: str) -> None:
    r = run._r
    fc_begin = OxmlElement("w:fldChar"); fc_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.text = field_name
    fc_end = OxmlElement("w:fldChar"); fc_end.set(qn("w:fldCharType"), "end")
    r.append(fc_begin); r.append(instr); r.append(fc_end)


def _inline(para, text: str, base_size: int = 11) -> None:
    pattern = re.compile(r"(\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`)")
    last = 0
    for m in pattern.finditer(text):
        if m.start() > last:
            r = para.add_run(text[last:m.start()])
            r.font.name = "Calibri"; r.font.size = Pt(base_size)
        if m.group(2):
            r = para.add_run(m.group(2)); r.bold = True; r.italic = True
            r.font.name = "Calibri"; r.font.size = Pt(base_size)
        elif m.group(3):
            r = para.add_run(m.group(3)); r.bold = True
            r.font.name = "Calibri"; r.font.size = Pt(base_size)
        elif m.group(4):
            r = para.add_run(m.group(4)); r.italic = True
            r.font.name = "Calibri"; r.font.size = Pt(base_size)
        elif m.group(5):
            r = para.add_run(m.group(5))
            r.font.name = "Courier New"; r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(0xC7, 0x25, 0x4F)
        last = m.end()
    if last < len(text):
        r = para.add_run(text[last:])
        r.font.name = "Calibri"; r.font.size = Pt(base_size)


def _apply_styles(doc: Document) -> None:
    def _style(name, base, font, size, bold=False, color=None, sb=0, sa=6):
        try:
            st = doc.styles[name]
        except KeyError:
            st = doc.styles.add_style(name, 1)
            try:
                st.base_style = doc.styles[base]
            except Exception:
                pass
        st.font.name = font
        st.font.size = Pt(size)
        st.font.bold = bold
        if color:
            st.font.color.rgb = RGBColor(*bytes.fromhex(color))
        st.paragraph_format.space_before = Pt(sb)
        st.paragraph_format.space_after = Pt(sa)

    _style("H1", "Heading 1", "Calibri", 18, True, "1A3C6B", 18, 6)
    _style("H2", "Heading 2", "Calibri", 14, True, "2563A8", 14, 4)
    _style("H3", "Heading 3", "Calibri", 12, True, "3B73C4", 10, 3)
    _style("H4", "Heading 4", "Calibri", 11, True, "555555", 8, 2)
    _style("Body", "Normal", "Calibri", 11, False, None, 0, 6)


def _build_cover_page(doc: Document, doc_type: str, project_name: str) -> None:
    doc.add_paragraph()

    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_before = Pt(0)
    cp.paragraph_format.space_after = Pt(0)
    if LOGO_PATH and os.path.exists(LOGO_PATH):
        cp.add_run().add_picture(LOGO_PATH, width=Inches(3.5))
    else:
        wr = cp.add_run(BRAND)
        wr.font.name = "Calibri"; wr.font.size = Pt(24); wr.bold = True
        wr.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6B)

    for _ in range(2):
        doc.add_paragraph()

    full_type = _DOC_FULLNAMES.get(doc_type, doc_type)
    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tp.paragraph_format.space_before = Pt(0)
    tp.paragraph_format.space_after = Pt(6)
    tr = tp.add_run(full_type)
    tr.font.name = "Calibri"; tr.font.size = Pt(36); tr.font.bold = True
    tr.font.color.rgb = RGBColor(0x0C, 0x23, 0x40)

    if project_name:
        pp = doc.add_paragraph()
        pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pp.paragraph_format.space_before = Pt(6)
        pp.paragraph_format.space_after = Pt(0)
        pr = pp.add_run(project_name)
        pr.font.name = "Calibri"; pr.font.size = Pt(22)
        pr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    for _ in range(2):
        doc.add_paragraph()

    mt = doc.add_table(rows=4, cols=2)
    mt.alignment = WD_TABLE_ALIGNMENT.CENTER
    mt.style = "Table Grid"
    mt.autofit = False
    try:
        mt.columns[0].width = Inches(2.0)
        mt.columns[1].width = Inches(4.0)
    except Exception:
        pass

    for i, (k, v) in enumerate([
        ("Prepared by", BRAND),
        ("Document Type", doc_type),
        ("Date", datetime.date.today().strftime("%B %d, %Y")),
        ("Status", "Draft"),
    ]):
        kc, vc = mt.rows[i].cells[0], mt.rows[i].cells[1]
        try:
            kc.width = Inches(2.0)
            vc.width = Inches(4.0)
        except Exception:
            pass
        _set_cell_bg(kc, "F0F4FA")
        kc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        vc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        kp = kc.paragraphs[0]
        kp.paragraph_format.space_before = Pt(3)
        kp.paragraph_format.space_after = Pt(3)
        kr = kp.add_run(k)
        kr.font.name = "Calibri"; kr.font.size = Pt(12); kr.bold = True
        kr.font.color.rgb = RGBColor(0x0C, 0x23, 0x40)
        vp = vc.paragraphs[0]
        vp.paragraph_format.space_before = Pt(3)
        vp.paragraph_format.space_after = Pt(3)
        vr = vp.add_run(v)
        vr.font.name = "Calibri"; vr.font.size = Pt(11)

    doc.add_page_break()


def _build_header_footer(doc: Document, doc_type: str) -> None:
    for sec in doc.sections:
        sec.different_first_page_header_footer = True
        hdr = sec.header
        hp = hdr.paragraphs[0] if hdr.paragraphs else hdr.add_paragraph()
        hp.clear()
        hp.paragraph_format.space_before = Pt(0)
        hp.paragraph_format.space_after = Pt(4)
        if LOGO_PATH and os.path.exists(LOGO_PATH):
            hp.add_run().add_picture(LOGO_PATH, height=Pt(20))
        br = hp.add_run(BRAND)
        br.font.name = "Calibri"; br.font.size = Pt(9); br.bold = True
        br.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6B)
        tr2 = hp.add_run(f"\t{doc_type}")
        tr2.font.name = "Calibri"; tr2.font.size = Pt(9); tr2.bold = True
        tr2.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6B)
        tabs = OxmlElement("w:tabs"); tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "right"); tab.set(qn("w:pos"), "9360")
        tabs.append(tab); hp._p.get_or_add_pPr().append(tabs)
        _add_border(hp)

        ftr = sec.footer
        fp = ftr.paragraphs[0] if ftr.paragraphs else ftr.add_paragraph()
        fp.clear()
        fp.paragraph_format.space_before = Pt(4)
        cr = fp.add_run(BRAND)
        cr.font.name = "Calibri"; cr.font.size = Pt(9)
        cr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        for txt, is_field in [
            ("\t\tPage ", False), ("PAGE", True), (" of ", False), ("NUMPAGES", True)
        ]:
            r = fp.add_run() if is_field else fp.add_run(txt)
            if is_field:
                _field_run(r, txt)
            r.font.name = "Calibri"; r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        pPr2 = fp._p.get_or_add_pPr()
        tabs2 = OxmlElement("w:tabs"); tab2 = OxmlElement("w:tab")
        tab2.set(qn("w:val"), "right"); tab2.set(qn("w:pos"), "9360")
        tabs2.append(tab2); pPr2.append(tabs2)
        _add_border(fp, "top", "CCCCCC", "4", "0")


def _render_markdown_body(doc: Document, md_text: str) -> None:
    lines = md_text.split("\n")
    in_code = False
    in_table = False
    code_lines = []
    table_lines = []

    def flush_table():
        nonlocal in_table, table_lines
        rows_data = [
            [c.strip() for c in tl.strip().strip("|").split("|")]
            for tl in table_lines
            if not re.match(r"^\|[-| :]+\|$", tl.strip())
        ]
        if not rows_data:
            in_table = False; table_lines = []; return
        max_cols = max(len(r) for r in rows_data)
        tbl = doc.add_table(rows=len(rows_data), cols=max_cols)
        tbl.style = "Table Grid"
        tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
        for ri, row_data in enumerate(rows_data):
            for ci in range(max_cols):
                cell = tbl.rows[ri].cells[ci]
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                if ri == 0:
                    _set_cell_bg(cell, "D9E6F5")
                elif ri % 2 == 0:
                    _set_cell_bg(cell, "F5F8FD")
                txt = row_data[ci].strip("*").strip() if ci < len(row_data) else ""
                p = cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                r = p.add_run(txt)
                r.font.name = "Calibri"; r.font.size = Pt(10)
                r.font.bold = (ri == 0)
                if ri == 0:
                    r.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6B)
                tcPr = cell._tc.get_or_add_tcPr()
                tcBdr = OxmlElement("w:tcBorders")
                for side in ("top", "left", "bottom", "right"):
                    bd = OxmlElement(f"w:{side}")
                    bd.set(qn("w:val"), "single")
                    bd.set(qn("w:sz"), "4")
                    bd.set(qn("w:color"), "CCCCCC")
                    tcBdr.append(bd)
                tcPr.append(tcBdr)
        doc.add_paragraph()
        in_table = False; table_lines = []

    for line in lines:
        s = line.rstrip()

        if s.startswith("```"):
            if in_table:
                flush_table()
            if in_code:
                if code_lines:
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.3)
                    r = p.add_run("\n".join(code_lines))
                    r.font.name = "Courier New"; r.font.size = Pt(9)
                code_lines = []; in_code = False
            else:
                in_code = True
            continue

        if in_code:
            code_lines.append(s); continue

        if s.startswith("|"):
            in_table = True; table_lines.append(s); continue
        elif in_table:
            flush_table()

        if re.match(r"^---+$|^\*\*\*+$|^___+$", s):
            p = doc.add_paragraph(); _add_border(p, "bottom", "CCCCCC", "4", "0"); continue
        if s.startswith("#### "):
            p = doc.add_paragraph(style="H4"); _inline(p, s[5:]); continue
        if s.startswith("### "):
            p = doc.add_paragraph(style="H3"); _inline(p, s[4:]); continue
        if s.startswith("## "):
            p = doc.add_paragraph(style="H2"); _inline(p, s[3:]); continue
        if s.startswith("# "):
            p = doc.add_paragraph(style="H1"); _inline(p, s[2:]); _add_border(p); continue
        if s.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.4)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            _add_border(p, "left", "2563A8", "12", "4"); _inline(p, s[2:]); continue

        bm = re.match(r"^(\s*)([-*+])\s+(.+)", s)
        if bm:
            level = min(len(bm.group(1)) // 2, 3)
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Inches(0.25 + level * 0.25)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            _inline(p, bm.group(3)); continue

        nm = re.match(r"^(\s*)(\d+)\.\s+(.+)", s)
        if nm:
            level = min(len(nm.group(1)) // 3, 3)
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.left_indent = Inches(0.25 + level * 0.25)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            _inline(p, nm.group(3)); continue

        if not s:
            continue
        p = doc.add_paragraph(style="Body"); _inline(p, s)

    if in_table:
        flush_table()


def build_docx(md_text: str, doc_type: str, project_name: str = "") -> io.BytesIO:
    """Render Markdown into a branded .docx and return it as an in-memory buffer."""
    logger.info("Building DOCX - doc_type=%s, length=%d chars", doc_type, len(md_text))
    doc = Document()
    for sec in doc.sections:
        sec.page_width = Inches(8.5)
        sec.page_height = Inches(11)
        sec.left_margin = sec.right_margin = sec.top_margin = sec.bottom_margin = Inches(1)

    _apply_styles(doc)
    _build_cover_page(doc, doc_type, project_name)
    _build_header_footer(doc, doc_type)
    _render_markdown_body(doc, md_text)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
